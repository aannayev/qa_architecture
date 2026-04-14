#!/usr/bin/env python3
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt


def add_code_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def strip_inline_code(text: str) -> str:
    return text.replace("`", "")


def convert(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    in_code_block = False

    for raw in lines:
        line = raw.rstrip("\n")

        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            add_code_paragraph(doc, line)
            continue

        if not line.strip():
            doc.add_paragraph("")
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            level = min(len(heading_match.group(1)), 6)
            text = strip_inline_code(heading_match.group(2).strip())
            doc.add_heading(text, level=level)
            continue

        bullet_match = re.match(r"^\s*-\s+(.*)$", line)
        if bullet_match:
            text = strip_inline_code(bullet_match.group(1).strip())
            doc.add_paragraph(text, style="List Bullet")
            continue

        number_match = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if number_match:
            text = strip_inline_code(number_match.group(1).strip())
            doc.add_paragraph(text, style="List Number")
            continue

        doc.add_paragraph(strip_inline_code(line))

    doc.save(str(docx_path))


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: md_to_docx.py <input.md> <output.docx>")
        return 1

    md_path = Path(sys.argv[1])
    docx_path = Path(sys.argv[2])

    if not md_path.exists():
        print(f"Input file not found: {md_path}")
        return 1

    convert(md_path, docx_path)
    print(f"Created: {docx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

